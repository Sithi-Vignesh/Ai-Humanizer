import io

import docx
from fastapi.responses import StreamingResponse
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer


def generate_docx(text: str) -> StreamingResponse:
    """Build a .docx document from *text* and return it as a StreamingResponse."""
    doc = docx.Document()
    doc.add_heading("Humanized Text", 0)
    for paragraph in text.split("\n"):
        if paragraph.strip():
            doc.add_paragraph(paragraph.strip())
    file_stream = io.BytesIO()
    doc.save(file_stream)
    file_stream.seek(0)
    return StreamingResponse(
        file_stream,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": "attachment; filename=humanized.docx"},
    )


def generate_pdf(text: str) -> StreamingResponse:
    """Build a PDF document from *text* and return it as a StreamingResponse."""
    file_stream = io.BytesIO()
    doc = SimpleDocTemplate(
        file_stream,
        pagesize=A4,
        rightMargin=72,
        leftMargin=72,
        topMargin=72,
        bottomMargin=18,
    )
    styles = getSampleStyleSheet()
    story = [Paragraph("Humanized Text", styles["Title"]), Spacer(1, 12)]
    for paragraph in text.split("\n"):
        if paragraph.strip():
            story.append(Paragraph(paragraph.strip(), styles["Normal"]))
            story.append(Spacer(1, 6))
    doc.build(story)
    file_stream.seek(0)
    return StreamingResponse(
        file_stream,
        media_type="application/pdf",
        headers={"Content-Disposition": "attachment; filename=humanized.pdf"},
    )
