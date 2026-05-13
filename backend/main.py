from fastapi import FastAPI, UploadFile, File, HTTPException
from fastapi.responses import StreamingResponse
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from pydantic import BaseModel
from fastapi.middleware.cors import CORSMiddleware
from groq import Groq
from dotenv import load_dotenv
from pathlib import Path
import os
import sys
import io
import fitz
import docx

from transformers import RobertaTokenizer, RobertaForSequenceClassification
import torch

# ── Model Path ─────────────────────────────────────────────
if getattr(sys, 'frozen', False):
    BASE_DIR = Path(sys.executable).parent
else:
    BASE_DIR = Path(__file__).parent

MODEL_PATH = BASE_DIR / "models" / "ai-detector-model-v3"

# ── Load RoBERTa Model ─────────────────────────────────────
tokenizer_local = RobertaTokenizer.from_pretrained(str(MODEL_PATH), local_files_only=True)
model_local = RobertaForSequenceClassification.from_pretrained(str(MODEL_PATH), local_files_only=True)
model_local.eval()
device = torch.device("cuda" if torch.cuda.is_available() else "cpu")
model_local.to(device)

load_dotenv()

# ── App Setup ──────────────────────────────────────────────
app = FastAPI()

app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_methods=["*"],
    allow_headers=["*"],
)

client = Groq(api_key=os.getenv("GROQ_API_KEY"))

class TextInput(BaseModel):
    text: str
    strength: str = "medium"

# ── Health ─────────────────────────────────────────────────
@app.get("/health")
def health():
    return {"status": "ok"}

# ── Detect ─────────────────────────────────────────────────
@app.post("/detect")
def detect(data: TextInput):
    if not data.text.strip():
        return {"label": "Unknown", "score": 0.0, "ai_percent": 0.0, "human_percent": 0.0}

    inputs = tokenizer_local(
        data.text[:512],
        return_tensors="pt",
        truncation=True,
        max_length=512
    ).to(device)

    with torch.no_grad():
        outputs = model_local(**inputs)
        probs = torch.softmax(outputs.logits, dim=1)
        ai_percent = probs[0][1].item()

    human_percent = 1.0 - ai_percent
    label = "AI" if ai_percent > 0.5 else "Human"
    score = ai_percent if label == "AI" else human_percent

    return {
        "label": label,
        "score": round(score, 4),
        "ai_percent": round(ai_percent * 100, 2),
        "human_percent": round(human_percent * 100, 2)
    }

# ── Humanize ───────────────────────────────────────────────
def chunk_by_paragraphs(text: str, max_chunk_size=1000):
    paragraphs = text.split('\n')
    chunks, current_chunk, current_len = [], [], 0
    for para in paragraphs:
        if current_len + len(para) <= max_chunk_size:
            current_chunk.append(para)
            current_len += len(para) + 1
        else:
            if current_chunk:
                chunks.append('\n'.join(current_chunk))
            current_chunk = [para]
            current_len = len(para) + 1
    if current_chunk:
        chunks.append('\n'.join(current_chunk))
    return chunks

@app.post("/humanize")
def humanize(data: TextInput):
    if not data.text.strip():
        return {"humanized": ""}

    instructions = {
        "light": "Keep the tone mostly similar to the original, just fix slight robotic phrasing, adjust grammar smoothly, and resolve disjointed sentences.",
        "medium": "Rewrite to sound completely natural and human. Use standard contractions, vary sentence length, and remove generic robotic phrasing.",
        "heavy": "Completely rewrite in a highly conversational, extremely organic style. Sound very human. Use slang where appropriate, strong contractions, and dynamic sentence structure."
    }
    strength_instruction = instructions.get(data.strength.lower(), instructions["medium"])
    prompt = f"You are a text rewriter. {strength_instruction} IMPORTANT RULES: Keep ALL content including names, titles, dates, numbers, headings, and metadata. Do NOT summarize, skip, condense, or remove any content. Rewrite every single line. Return ONLY the rewritten text, nothing else."

    chunks = chunk_by_paragraphs(data.text, max_chunk_size=1000)
    humanized_chunks = []

    for chunk in chunks:
        if not chunk.strip():
            humanized_chunks.append(chunk)
            continue
        try:
            response = client.chat.completions.create(
                model="llama-3.3-70b-versatile",
                messages=[
                    {"role": "system", "content": prompt},
                    {"role": "user", "content": chunk}
                ],
                temperature=0.3
            )
            humanized_chunks.append(response.choices[0].message.content.strip())
        except Exception as e:
            print(f"Error humanizing chunk: {e}")
            humanized_chunks.append(chunk)

    return {"humanized": "\n\n".join(humanized_chunks)}

# ── Extract ────────────────────────────────────────────────
@app.post("/extract")
async def extract(file: UploadFile = File(...)):
    filename = file.filename.lower() if file.filename else ""
    if not (filename.endswith('.pdf') or filename.endswith('.docx')):
        raise HTTPException(status_code=400, detail="Only .pdf and .docx files are supported")

    contents = await file.read()
    extracted_text = ""

    try:
        if filename.endswith('.pdf'):
            doc = fitz.open(stream=contents, filetype="pdf")
            for page in doc:
                extracted_text += page.get_text() + "\n"
            doc.close()
        elif filename.endswith('.docx'):
            doc = docx.Document(io.BytesIO(contents))
            for para in doc.paragraphs:
                extracted_text += para.text + "\n"
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Error extracting text: {str(e)}")

    return {"text": extracted_text.strip()}

# ── Export DOCX ────────────────────────────────────────────
@app.post("/export/docx")
def export_docx(data: TextInput):
    doc = docx.Document()
    doc.add_heading("Humanized Text", 0)
    for paragraph in data.text.split('\n'):
        if paragraph.strip():
            doc.add_paragraph(paragraph.strip())
    file_stream = io.BytesIO()
    doc.save(file_stream)
    file_stream.seek(0)
    return StreamingResponse(
        file_stream,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": "attachment; filename=humanized.docx"}
    )

# ── Export PDF ─────────────────────────────────────────────
@app.post("/export/pdf")
def export_pdf(data: TextInput):
    file_stream = io.BytesIO()
    doc = SimpleDocTemplate(file_stream, pagesize=A4,
                            rightMargin=72, leftMargin=72,
                            topMargin=72, bottomMargin=18)
    styles = getSampleStyleSheet()
    story = [Paragraph("Humanized Text", styles["Title"]), Spacer(1, 12)]
    for paragraph in data.text.split('\n'):
        if paragraph.strip():
            story.append(Paragraph(paragraph.strip(), styles["Normal"]))
            story.append(Spacer(1, 6))
    doc.build(story)
    file_stream.seek(0)
    return StreamingResponse(
        file_stream,
        media_type="application/pdf",
        headers={"Content-Disposition": "attachment; filename=humanized.pdf"}
    )

# ── Entry Point ────────────────────────────────────────────
if __name__ == "__main__":
    import uvicorn
    uvicorn.run(app, host="127.0.0.1", port=8000)